"""06: Genera DOSSIER.docx + carpeta organizada en RESEARCH/<slug>_<ts>/."""
from __future__ import annotations
import argparse, json, sys, shutil
from datetime import datetime
from pathlib import Path
from common import (RESEARCH_DEFAULT, dump_json, ensure_research_layout,
                     load_json, log_line, today_report, slugify_tesis)


def build_docx(triage_data: dict, output_path: Path, log) -> None:
    """Genera DOSSIER.docx ejecutivo."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()
    # Márgenes
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Style normal
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    tesis = triage_data.get("tesis", "")
    stats = triage_data.get("stats", {})
    califs = triage_data.get("calificaciones", [])

    # Portada
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("DOSSIER DE INVESTIGACIÓN JURISPRUDENCIAL")
    run.bold = True; run.font.size = Pt(16)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Comité de Jurisconsultos — análisis automatizado")
    run.italic = True; run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph(f"Fecha: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    doc.add_paragraph()

    # Sección 1 — Tesis
    doc.add_heading("1. Punto a acreditar (tesis)", level=1)
    p = doc.add_paragraph(tesis)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Sección 2 — Resumen
    doc.add_heading("2. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        f"Total de candidatos analizados: {sum(stats.get(k, 0) for k in ['KILL_SHOT','HIPERPERTINENTE','PERTINENTE','TANGENCIAL','NO_RELEVANTE'])}\n"
        f"  • KILL-SHOT: {stats.get('KILL_SHOT', 0)}\n"
        f"  • Hiperpertinentes: {stats.get('HIPERPERTINENTE', 0)}\n"
        f"  • Pertinentes: {stats.get('PERTINENTE', 0)}\n"
        f"  • Tangenciales (descartadas): {stats.get('TANGENCIAL', 0)}\n"
        f"  • No relevantes (descartadas): {stats.get('NO_RELEVANTE', 0)}\n"
        f"  • Adversas a la tesis: {stats.get('adversas', 0)}\n"
    )

    if stats.get("KILL_SHOT", 0) + stats.get("HIPERPERTINENTE", 0) == 0:
        p = doc.add_paragraph()
        run = p.add_run("⚠ ADVERTENCIA: 0 sentencias hiperpertinentes encontradas. "
                         "La tesis carece de cobertura jurisprudencial sólida en el "
                         "corpus consultado. Recomendaciones: (a) reformular con "
                         "instituciones más específicas, (b) ampliar período "
                         "temporal, (c) consultar doctrina académica.")
        run.font.color.rgb = RGBColor(0xCC, 0x44, 0x00)
        run.bold = True

    # Helper render por sentencia
    def render_sentencia(s: dict, level: int = 2):
        doc.add_heading(
            f"ROL {s.get('rol', '?')} — {s.get('tribunal', '?')}"
            f"{' (Sala ' + s['sala'] + ')' if s.get('sala') else ''}"
            f" — {s.get('fecha', '?')}",
            level=level)
        if s.get("caratulado"):
            p = doc.add_paragraph()
            run = p.add_run("Carátula: ")
            run.bold = True
            p.add_run(s["caratulado"])
        if s.get("adversa_a_tesis"):
            p = doc.add_paragraph()
            run = p.add_run("⚠ ADVERSA A LA TESIS")
            run.bold = True
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            if s.get("manejo_sugerido"):
                doc.add_paragraph(s["manejo_sugerido"])
        # Score
        sb = s.get("score_breakdown", {})
        if sb:
            doc.add_paragraph(
                f"Score: {s.get('score', 0):.2f} "
                f"(pronunciamiento {sb.get('pronunciamiento_density', 0):.2f}, "
                f"normas overlap {sb.get('normas_overlap', 0):.0%}, "
                f"jerarquía {sb.get('jerarquia_factor', 0):.1f})"
            )
        # Fundamento
        if s.get("fundamento_auto"):
            p = doc.add_paragraph()
            run = p.add_run("Fundamento: ")
            run.bold = True
            p.add_run(s["fundamento_auto"])
        # Cita verbatim
        cv = s.get("cita_verbatim") or {}
        if cv.get("texto"):
            p = doc.add_paragraph()
            run = p.add_run(f"Cita verbatim — Considerando {cv.get('considerando', '?')}°:")
            run.bold = True
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(f'"{cv["texto"]}"')
            run.italic = True
        # Normas
        if s.get("normas_centrales_detectadas"):
            doc.add_paragraph(
                f"Normas centrales detectadas: {', '.join(s['normas_centrales_detectadas'][:8])}"
            )
        if s.get("url_acceso"):
            doc.add_paragraph(f"URL pjud: {s['url_acceso']}")

    # Sección 3 — KILL-SHOT
    kill = [s for s in califs if s.get("calificacion") == "KILL_SHOT"]
    if kill:
        doc.add_heading("3. Autoridades centrales (KILL-SHOT)", level=1)
        for s in sorted(kill, key=lambda x: x.get("score", 0), reverse=True):
            render_sentencia(s, level=2)

    # Sección 4 — HIPERPERTINENTES
    hiper = [s for s in califs if s.get("calificacion") == "HIPERPERTINENTE"]
    if hiper:
        doc.add_heading("4. Sentencias hiperpertinentes", level=1)
        for s in sorted(hiper, key=lambda x: x.get("score", 0), reverse=True):
            render_sentencia(s, level=2)

    # Sección 5 — PERTINENTES (tabla compacta)
    perti = [s for s in califs if s.get("calificacion") == "PERTINENTE"]
    if perti:
        doc.add_heading("5. Sentencias pertinentes (tabla)", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "ROL"; hdr[1].text = "Tribunal"
        hdr[2].text = "Aporte"; hdr[3].text = "Cita verbatim"
        for s in sorted(perti, key=lambda x: x.get("score", 0), reverse=True):
            row = tbl.add_row().cells
            row[0].text = str(s.get("rol", ""))
            row[1].text = str(s.get("tribunal", ""))[:25]
            row[2].text = (s.get("fundamento_auto") or "")[:200]
            cv = s.get("cita_verbatim") or {}
            row[3].text = (cv.get("texto") or "")[:200]

    # Sección 6 — Apéndice descartadas
    descartadas = [s for s in califs
                    if s.get("calificacion") in ("TANGENCIAL", "NO_RELEVANTE", "FAILED")]
    if descartadas:
        doc.add_heading("6. Apéndice — sentencias descartadas", level=1)
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Light Grid Accent 1"
        hdr = tbl.rows[0].cells
        hdr[0].text = "ROL"; hdr[1].text = "Calificación"; hdr[2].text = "Razón de descarte"
        for s in descartadas[:50]:
            row = tbl.add_row().cells
            row[0].text = str(s.get("rol", ""))
            row[1].text = s.get("calificacion", "")
            row[2].text = s.get("razon_descarte", "")[:150]

    doc.save(output_path)
    log(f"  DOSSIER.docx → {output_path}")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--tesis", required=True)
    ap.add_argument("--root", type=Path, default=RESEARCH_DEFAULT)
    args = ap.parse_args()

    paths = ensure_research_layout(args.root)
    report = today_report(paths["reports"], "DOSSIER")
    log = lambda m: log_line(report, m)
    slug = slugify_tesis(args.tesis)

    log(f"=== 06 dossier === slug={slug}")
    triage_data = load_json(paths["work"] / f"TRIAGE_RESULT_{slug}.json", {})
    if not triage_data.get("calificaciones"):
        log("FALLA: TRIAGE_RESULT vacío o no existe")
        return 9

    # Crear carpeta destino
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    dest = paths["research"] / f"{slug}_{ts}"
    (dest / "PDFS").mkdir(parents=True, exist_ok=True)
    (dest / "MD").mkdir(parents=True, exist_ok=True)
    log(f"destino: {dest}")

    califs = triage_data["calificaciones"]
    incluidos = [s for s in califs if s.get("calificacion") in
                  ("KILL_SHOT", "HIPERPERTINENTE", "PERTINENTE")]
    log(f"incluidos en dossier: {len(incluidos)}")

    # Copiar MD-twins de incluidos a dest/MD
    copied_md = 0
    for s in incluidos:
        src = s.get("md_path")
        if src and Path(src).exists():
            try:
                rol_safe = (s.get("rol") or "unknown").replace("/", "-")
                shutil.copy2(src, dest / "MD" / f"{rol_safe}.md")
                copied_md += 1
            except Exception as e:
                log(f"  WARN copia MD {s.get('rol')}: {e}")
    log(f"  MD-twins copiados: {copied_md}")

    # DOSSIER.docx
    build_docx(triage_data, dest / "DOSSIER.docx", log)

    # INDEX.md
    idx_lines = [f"# INDEX — {args.tesis}", "", f"_Generado: {ts}_", ""]
    for cal in ["KILL_SHOT", "HIPERPERTINENTE", "PERTINENTE"]:
        items = [s for s in califs if s.get("calificacion") == cal]
        if items:
            idx_lines.append(f"## {cal} ({len(items)})")
            idx_lines.append("")
            idx_lines.append("| ROL | Tribunal | Score | Adversa | MD |")
            idx_lines.append("|---|---|---|---|---|")
            for s in sorted(items, key=lambda x: x.get("score", 0), reverse=True):
                rol_safe = (s.get("rol") or "?").replace("/", "-")
                idx_lines.append(
                    f"| {s.get('rol','')} | {s.get('tribunal','')[:25]} | "
                    f"{s.get('score', 0):.2f} | "
                    f"{'⚠' if s.get('adversa_a_tesis') else ''} | "
                    f"[MD](MD/{rol_safe}.md) |"
                )
            idx_lines.append("")
    (dest / "INDEX.md").write_text("\n".join(idx_lines), encoding="utf-8")
    log(f"  INDEX.md generado")

    # triage_log + manifest
    dump_json(dest / "triage_log.json", triage_data)
    dump_json(dest / "manifest.json", {
        "tesis": args.tesis, "slug": slug, "timestamp": ts,
        "total_candidatos": len(califs),
        "incluidos_dossier": len(incluidos),
        "stats": triage_data.get("stats"),
    })

    log(f"=== 06 OK === dossier en {dest}")
    print(f"\n✅ Dossier completo: {dest}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
